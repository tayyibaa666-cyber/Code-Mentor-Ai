from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Tuple
import io
import os
import time


# Public folder for storing quiz files
PUBLIC_QUIZ_FOLDER = "public/quizzes"
os.makedirs(PUBLIC_QUIZ_FOLDER, exist_ok=True)


def create_quiz_docx(quiz_data: Dict, session_id: str = "default", save_to_disk: bool = True) -> Tuple[bytes, str]:
    """
    Create a Word document (.docx) from quiz data.
    
    Args:
        quiz_data: Dictionary containing quiz_title and questions
        session_id: Session identifier for filename
        save_to_disk: Whether to save the file to disk
    
    Returns:
        Tuple[bytes, str]: The .docx file as bytes and file path (or None)
    """
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    
    heading_style = doc.styles['Heading 1']
    heading_font = heading_style.font
    heading_font.name = 'Arial'
    heading_font.size = Pt(14)
    heading_font.bold = True
    
    # Add title
    title = doc.add_heading(quiz_data.get('quiz_title', 'Programming Quiz'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add instructions
    doc.add_paragraph()
    instructions = doc.add_paragraph('Instructions: Please answer all questions. Select the best answer for each question.')
    instructions_format = instructions.runs[0].font
    instructions_format.italic = True
    instructions_format.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('-' * 80)
    doc.add_paragraph()
    
    # Add questions
    questions = quiz_data.get('questions', [])
    for i, q in enumerate(questions, 1):
        # Question number and text
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(f'Question {i}: {q.get("question", "")}')
        question_run.font.bold = True
        question_run.font.size = Pt(12)
        question_run.font.name = 'Arial'
        
        doc.add_paragraph()  # Spacing
        
        # Options
        options = q.get('options', [])
        labels = ['A', 'B', 'C', 'D']
        for j, option in enumerate(options):
            if j < len(labels):
                option_para = doc.add_paragraph(f'{labels[j]}. {option}', style='List Bullet')
                option_para.paragraph_format.left_indent = Inches(0.5)
                option_para.runs[0].font.size = Pt(11)
                option_para.runs[0].font.name = 'Arial'
        
        doc.add_paragraph()  # Spacing
        
        # Answer section (for teacher/answer key)
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run(f'Correct Answer: {q.get("correct_answer", "N/A")}')
        answer_run.font.bold = True
        answer_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
        answer_run.font.size = Pt(10)
        
        # Explanation
        explanation = q.get('explanation', '')
        if explanation:
            explanation_para = doc.add_paragraph()
            explanation_run = explanation_para.add_run(f'Explanation: {explanation}')
            explanation_run.font.italic = True
            explanation_run.font.size = Pt(10)
            explanation_run.font.color.rgb = RGBColor(100, 100, 100)  # Gray color
        
        doc.add_paragraph()
        doc.add_paragraph('-' * 80)
        doc.add_paragraph()
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    file_bytes = buffer.getvalue()
    
    # Save to disk if requested
    file_path = None
    if save_to_disk:
        timestamp = int(time.time())
        filename = f"quiz_answers_{session_id}_{timestamp}.docx"
        file_path = os.path.join(PUBLIC_QUIZ_FOLDER, filename)
        with open(file_path, 'wb') as f:
            f.write(file_bytes)
    
    return file_bytes, file_path


def create_quiz_docx_student_version(quiz_data: Dict, session_id: str = "default", save_to_disk: bool = True) -> Tuple[bytes, str]:
    """
    Create a student version of the quiz (without answers).
    
    Args:
        quiz_data: Dictionary containing quiz_title and questions
        session_id: Session identifier for filename
        save_to_disk: Whether to save the file to disk
    
    Returns:
        Tuple[bytes, str]: The .docx file as bytes and file path (or None)
    """
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    title_font.bold = True
    
    # Add title
    title = doc.add_heading(quiz_data.get('quiz_title', 'Programming Quiz'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add instructions
    doc.add_paragraph()
    instructions = doc.add_paragraph('Instructions: Please answer all questions. Select the best answer for each question.')
    instructions_format = instructions.runs[0].font
    instructions_format.italic = True
    instructions_format.size = Pt(11)
    
    doc.add_paragraph()
    doc.add_paragraph('Name: _______________________')
    doc.add_paragraph('Date: _______________________')
    doc.add_paragraph()
    doc.add_paragraph('-' * 80)
    doc.add_paragraph()
    
    # Add questions
    questions = quiz_data.get('questions', [])
    for i, q in enumerate(questions, 1):
        # Question number and text
        question_para = doc.add_paragraph()
        question_run = question_para.add_run(f'Question {i}: {q.get("question", "")}')
        question_run.font.bold = True
        question_run.font.size = Pt(12)
        question_run.font.name = 'Arial'
        
        doc.add_paragraph()  # Spacing
        
        # Options
        options = q.get('options', [])
        labels = ['A', 'B', 'C', 'D']
        for j, option in enumerate(options):
            if j < len(labels):
                option_para = doc.add_paragraph(f'{labels[j]}. {option}', style='List Bullet')
                option_para.paragraph_format.left_indent = Inches(0.5)
                option_para.runs[0].font.size = Pt(11)
                option_para.runs[0].font.name = 'Arial'
        
        doc.add_paragraph()  # Spacing
        doc.add_paragraph('-' * 80)
        doc.add_paragraph()
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    file_bytes = buffer.getvalue()
    
    # Save to disk if requested
    file_path = None
    if save_to_disk:
        timestamp = int(time.time())
        filename = f"quiz_{session_id}_{timestamp}.docx"
        file_path = os.path.join(PUBLIC_QUIZ_FOLDER, filename)
        with open(file_path, 'wb') as f:
            f.write(file_bytes)
    
    return file_bytes, file_path

